"""Document upload flow: storage upload + Claude summary + chunk + insert."""

from __future__ import annotations

import mimetypes
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Optional
from uuid import UUID, uuid4

from goldman.chunker import chunk_text


_SAFE_NAME = re.compile(r"[^A-Za-z0-9._-]")


def _safe_filename(name: str) -> str:
    return _SAFE_NAME.sub("_", name)


def extract_text_from_pdf(file_path: Path) -> str:
    """Pull raw text from a PDF using pypdf."""
    from pypdf import PdfReader
    reader = PdfReader(str(file_path))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    return "\n\n".join(pages).strip()


def _sniff_mime(file_path: Path) -> Optional[str]:
    """Detect the real file type from its magic bytes, ignoring the filename.

    Telegram photos arrive with no filename, and our intake sometimes names
    them 'document.pdf' — so an extension is not trustworthy. Reading the
    first few bytes is. Returns a mime string for the types we know how to
    handle, or None when we can't tell (caller falls back to the extension).
    """
    try:
        head = file_path.read_bytes()[:8]
    except Exception:
        return None
    if head[:3] == b"\xff\xd8\xff":
        return "image/jpeg"
    if head[:8] == b"\x89PNG\r\n\x1a\n":
        return "image/png"
    if head[:4] == b"%PDF":
        return "application/pdf"
    if head[:2] == b"PK":
        # docx/xlsx are zip containers; defer to the extension to disambiguate.
        return None
    return None


def extract_text_from_docx(file_path: Path) -> str:
    """Pull paragraphs + table cell text out of a .docx."""
    from docx import Document
    doc = Document(str(file_path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_text_from_xlsx(file_path: Path) -> str:
    """Pull cell values out of every sheet as pipe-separated rows."""
    from openpyxl import load_workbook
    wb = load_workbook(str(file_path), read_only=True, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"## Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
        parts.append("")
    return "\n".join(parts).strip()


VISION_OCR_MIN_CHARS = 50  # threshold below which we fall back to vision OCR


def _read_text(file_path: Path, mime_type: str) -> str:
    # Trust the file's actual bytes over its (possibly wrong) extension.
    # A screenshot sent through Telegram can reach here named 'document.pdf';
    # without this, pypdf would choke on JPEG data and crash the whole reply.
    sniffed = _sniff_mime(file_path)
    if sniffed:
        mime_type = sniffed
    if mime_type == "application/pdf":
        from goldman.llm import vision_extract_text
        try:
            text = extract_text_from_pdf(file_path)
        except Exception:
            # Corrupt / truncated / not-really-a-PDF — OCR it instead of dying.
            text = ""
        if len(text.strip()) < VISION_OCR_MIN_CHARS:
            # Image-only / scanned / unreadable PDF — fall back to vision OCR.
            try:
                text = vision_extract_text(file_path=file_path) or text
            except Exception:
                pass
        return text
    suffix = file_path.suffix.lower()
    if suffix == ".docx" or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_text_from_docx(file_path)
    if suffix == ".xlsx" or mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        return extract_text_from_xlsx(file_path)
    # Image files fall straight through to vision OCR.
    if mime_type.startswith("image/"):
        from goldman.llm import vision_extract_text
        return vision_extract_text(file_path=file_path)
    return file_path.read_text(errors="replace")


@dataclass
class UploadResult:
    document_id: UUID
    chunk_count: int
    storage_path: str


def upload_document(
    *,
    file_path: Path,
    entity_id: Optional[UUID],
    entity_slug: Optional[str],
    storage,
    doc_repo,
    chunk_repo,
    summariser,
    bucket: str,
    source: str = "uploaded",
    chunk_max_tokens: int = 512,
    chunk_overlap_tokens: int = 64,
    pack_topic: Optional[str] = None,
    pack_version: Optional[str] = None,
    storage_path_override: Optional[str] = None,
    drive_client=None,
    drive_root_id: Optional[str] = None,
    entity_legal_name: Optional[str] = None,
    drive_category: str = "Documents",
) -> UploadResult:
    """Upload one document end-to-end.

    storage      — SupabaseStorage instance
    doc_repo     — DocumentRepository instance
    chunk_repo   — DocumentChunkRepository instance
    summariser   — anything with .summarise(text) -> str

    For knowledge packs: pass storage_path_override (e.g.
    'packs/{topic}/{version}/...'), source='knowledge_pack', pack_topic,
    pack_version, and entity_id=entity_slug=None.
    """
    mime_type, _ = mimetypes.guess_type(file_path.name)
    mime_type = mime_type or "application/octet-stream"

    # 1. Storage upload
    body = file_path.read_bytes()
    if storage_path_override is not None:
        storage_path = storage_path_override
    else:
        year = datetime.utcnow().year
        storage_path = f"{entity_slug}/{year}/{uuid4().hex[:8]}-{_safe_filename(file_path.name)}"
    storage.upload(
        bucket=bucket,
        path=storage_path,
        content=body,
        content_type=mime_type,
    )

    # 2. Insert metadata row
    doc_id = doc_repo.insert(
        entity_id=entity_id,
        filename=file_path.name,
        mime_type=mime_type,
        source=source,
        original_storage_path=storage_path,
        pack_topic=pack_topic,
        pack_version=pack_version,
    )

    # 3. Extract text and chunk
    text = _read_text(file_path, mime_type)
    chunks = chunk_text(
        text,
        max_tokens=chunk_max_tokens,
        overlap_tokens=chunk_overlap_tokens,
    )

    # 4. Insert chunks
    for idx, chunk in enumerate(chunks):
        chunk_repo.insert(
            document_id=doc_id,
            chunk_index=idx,
            text=chunk,
        )

    # 5. Summarise (one-shot via Claude Haiku)
    if text.strip():
        try:
            summary = summariser.summarise(text)
            doc_repo.set_summary(doc_id, summary)
        except Exception:
            # Non-fatal: chunks are searchable even without a summary.
            pass

    # 6. Mirror to Google Drive (best effort; failure here is non-fatal).
    if drive_client is not None and entity_legal_name:
        try:
            from goldman.drive.folders import ensure_path
            year = str(datetime.utcnow().year)
            folder_id = ensure_path(
                drive_client,
                [entity_legal_name, year, drive_category],
                root_id=drive_root_id,
            )
            drive_client.upload_file(
                name=file_path.name,
                parent_id=folder_id,
                content=body,
                mime_type=mime_type,
            )
        except Exception:
            # Drive failure shouldn't block memory ingestion. We can sweep
            # missing-Drive docs later via a backfill if needed.
            pass

    return UploadResult(
        document_id=doc_id,
        chunk_count=len(chunks),
        storage_path=storage_path,
    )
